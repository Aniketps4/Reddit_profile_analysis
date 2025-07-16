import json
import argparse
from docx import Document
from collections import defaultdict

def analyze_text_for_themes(texts):
    """Simulate LLM to extract key themes and generate a custom quote."""
    themes = defaultdict(int)
    theme_keywords = {
        "business": ["business", "work", "professional"],
        "location": ["lucknow", "lko", "delhi", "city"],
        "health": ["healthy", "cook", "diet", "meal", "nutrition"],
        "law": ["cop", "fine", "bribe", "law"],
        "community": ["discuss", "lucknow", "delhi"]
    }
    
    all_text = " ".join(t.get("body", t.get("title", "")) for t in texts if isinstance(t.get("body", t.get("title", "")), str)).lower()
    for theme, keywords in theme_keywords.items():
        if any(k in all_text for k in keywords):
            themes[theme] += 1
    
    dominant_theme = max(themes, key=themes.get, default="health")
    quotes = {
        "business": "I aim to grow my career and business opportunities.",
        "location": "I want to explore and adapt to my new surroundings.",
        "health": "I want to improve my health and dietary habits.",
        "law": "I seek fair treatment in my daily life.",
        "community": "I enjoy connecting with my community."
    }
    return quotes.get(dominant_theme, "I want to improve my life experience.")

def build_persona(data):
    """Build a user persona based on scraped data."""
    persona = {
        "citations": {}
    }

    # Basic Info
    persona["name"] = data.get("username", "Unknown")
    persona["age"] = "Unknown"
    persona["occupation"] = "Unknown"
    persona["location"] = "Unknown"
    persona["tier"] = "Unknown"
    persona["archetype"] = "Unknown"
    for post in data["posts"]:
        if "body" in post and post["body"] and isinstance(post["body"], str):
            text = post["body"].lower()
            if "lucknow" in text or "lko" in text:
                persona["location"] = "Lucknow, India"
                persona["citations"]["location"] = post["url"]
            if "business" in text or "work" in text:
                persona["occupation"] = "Possible business professional"
                persona["citations"]["occupation"] = post["url"]
            if "new" in text or "shift" in text:
                persona["tier"] = "Early Adopters"
                persona["archetype"] = "The Explorer"
                persona["citations"]["tier_archetype"] = post["url"]

    # Motivations
    motivations = {"convenience": 0, "wellness": 0, "speed": 0, "preferences": 0, "comfort": 0, "dietary needs": 0}
    keywords = {
        "convenience": ["easy", "quick"],
        "wellness": ["healthy", "well", "power"],
        "speed": ["fast", "quick"],
        "preferences": ["like", "prefer", "love"],
        "comfort": ["relax", "enjoy", "cozy", "power"],
        "dietary needs": ["diet", "vegan", "eat", "food", "meal", "cook", "menu", "nutrition"]
    }
    all_text = [post["body"] for post in data["posts"] if "body" in post] + [c["body"] for c in data["comments"] if "body" in c]
    for text in all_text:
        if text and isinstance(text, str):
            text = text.lower()
            for mot, kw in keywords.items():
                if any(k in text for k in kw):
                    motivations[mot] += 1
    persona["motivations"] = {k: min(v, 5) for k, v in motivations.items()}

    # Personality
    persona["personality"] = {
        "introvert/extrovert": "Neutral",
        "intuition/sensing": "Neutral",
        "feeling/thinking": "Neutral",
        "perceiving/judging": "Neutral"
    }
    for comment in data["comments"]:
        text = comment.get("body", "").lower()
        if "😂" in text or "great" in text or "awesome" in text:
            persona["personality"]["feeling/thinking"] = "Feeling"
        if "new" in text or "shift" in text or "future" in text:
            persona["personality"]["intuition/sensing"] = "Intuition"

    # Behavior & Habits
    habits = []
    for comment in data["comments"]:
        text = comment.get("body", "").lower()
        if ("cook" in text or "menu" in text) and "habits_cook" not in persona["citations"]:
            habits.append("Explores healthy cooking habits")
            persona["citations"]["habits_cook"] = comment["url"]
        if ("lucknow" in text or "delhi" in text or "community" in text) and "habits_community" not in persona["citations"]:
            habits.append("Engages in community discussions")
            persona["citations"]["habits_community"] = comment["url"]
    persona["behavior_habits"] = habits if habits else ["No clear habits identified"]

    # Frustrations
    frustrations = []
    for comment in data["comments"]:
        text = comment.get("body", "").lower()
        if ("cop" in text or "fine" in text or "bribe" in text) and "frustrations_law" not in persona["citations"]:
            frustrations.append("Expresses frustration with law enforcement")
            persona["citations"]["frustrations_law"] = comment["url"]
    persona["frustrations"] = frustrations if frustrations else ["No clear frustrations identified"]

    # Goals & Needs
    goals = []
    for post in data["posts"]:
        text = post.get("body", "").lower()
        if ("cook" in text or "healthy" in text) and "goals_diet" not in persona["citations"]:
            goals.append("Aims to improve dietary habits")
            persona["citations"]["goals_diet"] = post["url"]
        if ("lucknow" in text or "shift" in text) and "goals_adapt" not in persona["citations"]:
            goals.append("Seeks to adapt to new environment")
            persona["citations"]["goals_adapt"] = post["url"]
    for comment in data["comments"]:
        text = comment.get("body", "").lower()
        if ("cook" in text or "healthy" in text) and "goals_diet" not in persona["citations"]:
            goals.append("Aims to improve dietary habits")
            persona["citations"]["goals_diet"] = comment["url"]
        if ("lucknow" in text or "shift" in text) and "goals_adapt" not in persona["citations"]:
            goals.append("Seeks to adapt to new environment")
            persona["citations"]["goals_adapt"] = comment["url"]
    persona["goals_needs"] = goals if goals else ["No clear goals identified"]

    # Custom Quote using LLM simulation
    all_content = data["posts"] + data["comments"]
    persona["quote"] = analyze_text_for_themes(all_content)

    return persona

def generate_word_file(persona, username):
    """Generate a Word document with the persona."""
    doc = Document()
    
    # Basic Info
    doc.add_heading("Persona for " + persona["name"], level=1)
    doc.add_paragraph(f"Age: {persona['age']}")
    doc.add_paragraph(f"Occupation: {persona['occupation']} (Source: {persona['citations'].get('occupation', 'N/A')})")
    doc.add_paragraph(f"Location: {persona['location']} (Source: {persona['citations'].get('location', 'N/A')})")
    doc.add_paragraph(f"Tier: {persona['tier']} (Source: {persona['citations'].get('tier_archetype', 'N/A')})")
    doc.add_paragraph(f"Archetype: {persona['archetype']} (Source: {persona['citations'].get('tier_archetype', 'N/A')})")

    # Motivations
    doc.add_heading("Motivations", level=2)
    for mot, score in persona["motivations"].items():
        doc.add_paragraph(f"- {mot}: {'█' * score} (Score: {score}/5)")

    # Personality
    doc.add_heading("Personality", level=2)
    for trait, value in persona["personality"].items():
        doc.add_paragraph(f"- {trait.replace('/', ' / ')}: {'█' * (5 if value == 'Neutral' else 3)} {value}")

    # Behavior & Habits
    doc.add_heading("Behavior & Habits", level=2)
    for habit in persona["behavior_habits"]:
        if habit == "Explores healthy cooking habits":
            source_key = "habits_cook"
        elif habit == "Engages in community discussions":
            source_key = "habits_community"
        else:
            source_key = "N/A"
        doc.add_paragraph(f"- {habit} (Source: {persona['citations'].get(source_key, 'N/A')})")

    # Frustrations
    doc.add_heading("Frustrations", level=2)
    for frust in persona["frustrations"]:
        if frust == "Expresses frustration with law enforcement":
            source_key = "frustrations_law"
        doc.add_paragraph(f"- {frust} (Source: {persona['citations'].get(source_key, 'N/A')})")

    # Goals & Needs
    doc.add_heading("Goals & Needs", level=2)
    for goal in persona["goals_needs"]:
        if goal == "Aims to improve dietary habits":
            source_key = "goals_diet"
        elif goal == "Seeks to adapt to new environment":
            source_key = "goals_adapt"
        doc.add_paragraph(f"- {goal} (Source: {persona['citations'].get(source_key, 'N/A')})")

    # Quote
    doc.add_heading("Quote", level=2)
    doc.add_paragraph(persona["quote"])

    # Save the document
    doc.save(f"{username}_persona.docx")
    print(f"Word document saved as {username}_persona.docx")

def main():
    # Parse command line argument for JSON data file
    parser = argparse.ArgumentParser(description="Generate a Word document from scraped Reddit data.")
    parser.add_argument("json_file", help="Path to the JSON data file (e.g., Hungry-Move-6603_data.json)")
    args = parser.parse_args()
    
    # Load data from JSON file
    with open(args.json_file, "r", encoding="utf-8") as f:
        data = json.load(f)
    
    # Extract username from filename
    username = args.json_file.split("_data.json")[0]
    data["username"] = username
    
    # Build persona
    print(f"Building persona for {username}...")
    persona = build_persona(data)
    
    # Generate Word file
    print(f"Generating Word document for {username}...")
    generate_word_file(persona, username)
    print("Done!")

if __name__ == "__main__":
    main() 